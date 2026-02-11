import hashlib
import re
from pathlib import Path
from typing import List, Optional, Dict

from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn

from docs2vecs.subcommands.indexer.config.config import Config
from docs2vecs.subcommands.indexer.document.chunk import Chunk
from docs2vecs.subcommands.indexer.document.document import Document
from docs2vecs.subcommands.indexer.skills.skill import IndexerSkill


class ConfluenceFAQSplitter(IndexerSkill):
    """
    Advanced Q&A extractor for DOCX files with sophisticated parsing.
    
    Features:
    - Extracts Q&A pairs from FAQ DOCX documents
    - Ignores Table of Contents and Summary sections
    - Handles heading-based questions with '?' or problem patterns
    - Extracts hyperlinks from answers
    - Stops at 'Related articles' sections
    - Preserves table content in answers
    - Each Q&A pair becomes a single atomic chunk for optimal RAG retrieval
    
    Configuration parameters (all optional with sensible defaults):
    - min_heading_level: Minimum heading level for questions (default: 2)
    - max_heading_level: Maximum heading level for questions (default: 6)
    - skip_patterns: List of text patterns to skip in answer content (default: ['CONFIDENTIAL', 'Search the FAQ', 'Search Artifactory FAQ'])
    - skip_headings: List of heading titles to skip as questions (default: ['summary'])
    - question_patterns: List of prefixes that indicate a question/problem statement (default: ['i am ', 'i cannot ', ...])
    - stop_sections: List of regex patterns for sections that end Q&A extraction (default: ['related articles', 'see also'])
    """

    # Default configuration values
    DEFAULT_MIN_HEADING_LEVEL = 2
    DEFAULT_MAX_HEADING_LEVEL = 6
    DEFAULT_SKIP_PATTERNS = ['CONFIDENTIAL', 'Search the FAQ', 'Search Artifactory FAQ']
    DEFAULT_SKIP_HEADINGS = ['summary']
    DEFAULT_QUESTION_PATTERNS = [
        'i am ', 'i cannot ', "i can't ", 'i see ',
        'i have ', 'i need ', 'my ', 'when i ',
        'how do i ', 'how can i ', 'what is ', 'what are ',
        'why does ', 'why is ', 'where is ', 'where can '
    ]
    DEFAULT_STOP_SECTIONS = [
        r'^\s*related\s*articles?\s*$',
        r'^\s*related\s*resources?\s*$',
        r'^\s*see\s*also\s*$'
    ]

    def __init__(self, config: dict, global_config: Config):
        super().__init__(config, global_config)
        
        # Load configurable parameters with defaults
        self.min_heading_level = self._config.get('min_heading_level', self.DEFAULT_MIN_HEADING_LEVEL)
        self.max_heading_level = self._config.get('max_heading_level', self.DEFAULT_MAX_HEADING_LEVEL)
        self.skip_patterns = self._config.get('skip_patterns', self.DEFAULT_SKIP_PATTERNS)
        self.skip_headings = [h.lower() for h in self._config.get('skip_headings', self.DEFAULT_SKIP_HEADINGS)]
        self.question_patterns = [p.lower() for p in self._config.get('question_patterns', self.DEFAULT_QUESTION_PATTERNS)]
        
        # Compile stop section regexes
        stop_sections = self._config.get('stop_sections', self.DEFAULT_STOP_SECTIONS)
        self.related_res = [re.compile(p, re.I) for p in stop_sections]
        
        self.dot_leader_re = re.compile(r"\.{2,}\s*\d{1,4}\s*$")
        self.page_number_trail_re = re.compile(r"\s\d{1,4}\s*$")
        # Regex to match markdown-style links: [Link](URL)
        self.markdown_link_re = re.compile(r'\[Link\]\([^\)]+\)')
        
        self.logger.debug(f"ConfluenceFAQSplitter config: heading_levels={self.min_heading_level}-{self.max_heading_level}, "
                          f"skip_patterns={len(self.skip_patterns)}, question_patterns={len(self.question_patterns)}")

    def run(self, input: Optional[List[Document]] = None) -> List[Document]:
        self.logger.info("Running ConfluenceFAQSplitter...")

        if not input:
            self.logger.error("No documents provided in input")
            return []

        for doc in input:
            self.logger.debug(f"Processing document: {doc.filename}")
            
            # Check if file is a DOCX
            filename_str = str(doc.filename)
            if not filename_str.lower().endswith('.docx'):
                self.logger.warning(f"Skipping non-DOCX file: {doc.filename}")
                continue
            
            try:
                qa_pairs = self._extract_qa_from_docx(doc.filename)
                self.logger.info(f"Extracted {len(qa_pairs)} Q&A pairs from {doc.filename}")

                for idx, qa_data in enumerate(qa_pairs, 1):
                    question = qa_data['question']
                    answer = qa_data['answer']

                    if not question.strip() or not answer.strip():
                        self.logger.debug(f"Skipping Q&A pair {idx} - missing question or answer")
                        continue

                    links = qa_data.get('links', [])
                    
                    # Filter out links where the text is just the URL itself (redundant)
                    # Only include links with meaningful descriptive text in the References section
                    meaningful_links = [
                        link for link in links 
                        if not self._is_link_text_redundant(link['text'], link['url'])
                    ]
                    
                    # Format links for inclusion in content
                    links_text = ""
                    if meaningful_links:
                        links_list = [f"- {link['text']}: {link['url']}" for link in meaningful_links]
                        links_text = f"\n\nReferences (hyperlinks from the answer):\n" + "\n".join(links_list)
                    
                    # Combine question and answer into a single chunk
                    combined_text = f"Q: {question}\n\nA: {answer}{links_text}"
                    
                    chunk = Chunk()
                    chunk.document_id = hashlib.sha256(combined_text.encode()).hexdigest()
                    chunk.document_name = Path(doc.filename).name
                    chunk.tag = doc.tag
                    chunk.content = combined_text  # Full Q&A for retrieval
                    chunk.chunk_id = f"{chunk.document_id}_{idx}"
                    chunk.source_link = doc.source_url or ""
                    
                    doc.add_chunk(chunk)

                self.logger.debug(f"Split {doc.filename} into {len(doc.chunks)} Q&A chunks")

            except Exception as e:
                self.logger.error(f"Error processing {doc.filename}: {e}", exc_info=True)
                continue

        return input

    def _extract_qa_from_docx(self, docx_path: str) -> List[Dict[str, any]]:
        """
        Extract Q&A pairs from a DOCX file with sophisticated parsing.
        
        Returns:
            List of dicts with 'question', 'answer', 'links' keys
        """
        doc = DocxDocument(docx_path)
        qa: List[Dict[str, any]] = []

        current_q = None
        current_ans: List[str] = []
        current_links: List[Dict[str, str]] = []
        in_toc = False
        in_summary = False

        for blk in self._iter_block_items(doc):
            text = self._norm(self._block_text(blk))
            lvl = self._heading_level(blk) if isinstance(blk, Paragraph) else None
            title = self._norm(blk.text).lower() if isinstance(blk, Paragraph) else ""

            # ---------- TOC detection ----------
            if isinstance(blk, Paragraph) and title in ('table of contents', 'contents'):
                in_toc = True
                # Finalize any open Q before entering TOC
                if current_q is not None:
                    qa.append({
                        "question": current_q,
                        "answer": "\n".join(current_ans).strip() if current_ans else "",
                        "links": current_links,
                    })
                current_q, current_ans, current_links = None, [], []
                continue

            # End TOC at next major heading (H1 or H2)
            if in_toc and isinstance(blk, Paragraph) and lvl is not None and lvl <= 2 and title not in ('table of contents', 'contents'):
                in_toc = False

            # Skip TOC-styled paragraphs and dotted leader lines
            if in_toc and isinstance(blk, Paragraph):
                if self._style_name(blk).lower().startswith('toc') or self._is_toc_line(text):
                    continue

            # ---------- Skip-heading regions (e.g. Summary, Overview) ----------
            if isinstance(blk, Paragraph) and lvl is not None and title in self.skip_headings:
                in_summary = True
                continue
            
            if in_summary and isinstance(blk, Paragraph) and lvl is not None and title not in self.skip_headings:
                in_summary = False

            # ---------- Stop at Related articles ----------
            if current_q is not None and self._is_related_heading(blk):
                qa.append({
                    "question": current_q,
                    "answer": "\n".join(current_ans).strip() if current_ans else "",
                    "links": current_links,
                })
                current_q, current_ans, current_links = None, [], []
                continue

            # ---------- New question ----------
            if not in_toc and not in_summary and self._is_question_block(blk):
                # Flush previous
                if current_q is not None:
                    qa.append({
                        "question": current_q,
                        "answer": "\n".join(current_ans).strip() if current_ans else "",
                        "links": current_links,
                    })
                current_q = self._norm(blk.text)
                current_ans = []
                current_links = []
                continue

            # ---------- Accumulate answer ----------
            if current_q is not None:
                # Skip banners/boilerplate using configurable patterns
                if any(text.upper() == pat.upper() for pat in self.skip_patterns):
                    continue
                if text:
                    block_text = self._block_text(blk)
                    # Remove markdown-style links [Link](URL)
                    cleaned_text = self._remove_markdown_links(block_text)
                    if cleaned_text:  # Only add if there's content left after cleaning
                        current_ans.append(cleaned_text)
                    # Extract hyperlinks from block
                    links = self._extract_hyperlinks_from_block(blk)
                    current_links.extend(links)

        # Finalize last Q
        if current_q is not None:
            qa.append({
                "question": current_q,
                "answer": "\n".join(current_ans).strip() if current_ans else "",
                "links": current_links,
            })

        return qa

    # ---------- Helper methods ----------
    
    def _iter_block_items(self, parent):
        """Yield paragraphs and tables in document order."""
        body = parent.element.body
        for child in body.iterchildren():
            if child.tag == qn('w:p'):
                yield Paragraph(child, parent)
            elif child.tag == qn('w:tbl'):
                yield Table(child, parent)

    def _extract_table_text(self, table) -> str:
        """Recursively extract text from a table, including nested tables."""
        parts = []
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text or '')
                for nested_table in cell.tables:
                    nested_text = self._extract_table_text(nested_table)
                    if nested_text:
                        parts.append(nested_text)
        
        out = []
        for t in (x.strip() for x in parts):
            if t == '' and (not out or out[-1] == ''):
                continue
            out.append(t)
        return "\n".join(out).strip()

    def _block_text(self, block) -> str:
        """Get text from a paragraph or table."""
        if isinstance(block, Paragraph):
            return block.text or ''
        if isinstance(block, Table):
            return self._extract_table_text(block)
        return ''

    def _style_name(self, par: Paragraph) -> str:
        try:
            return (par.style.name or '').strip()
        except Exception:
            return ''

    def _heading_level(self, par: Paragraph):
        if not isinstance(par, Paragraph):
            return None
        m = re.match(r'Heading\s*(\d+)$', self._style_name(par))
        return int(m.group(1)) if m else None

    def _norm(self, s: str) -> str:
        """Normalize text by replacing non-breaking spaces and collapsing whitespace."""
        s = (s or '').replace('\u00A0', ' ')
        s = re.sub(r'\s{2,}', ' ', s.strip())
        return s
    
    def _remove_markdown_links(self, s: str) -> str:
        """Remove markdown-style links in the form [Link](URL) from text."""
        return self.markdown_link_re.sub('', s).strip()
    
    def _is_link_text_redundant(self, text: str, url: str) -> bool:
        """
        Check if hyperlink text is redundant (i.e., it's just the URL itself or very similar).
        References:
        - https://example.com: https://example.com   ← Useless, filtered out
        
        Returns True if the link text is redundant and should be excluded from the References section.
        This keeps the answer clean while preserving meaningful link descriptions.
        """
        # Normalize both for comparison
        text_normalized = text.strip().lower()
        url_normalized = url.strip().lower()
        
        # Strip common trailing characters that might be added accidentally
        # (parentheses, periods, commas, etc.)
        text_cleaned = text_normalized.rstrip(').,;: ')
        url_cleaned = url_normalized.rstrip(').,;: ')
        
        # Remove common URL prefixes for comparison
        url_without_protocol = re.sub(r'^https?://', '', url_cleaned)
        url_without_www = re.sub(r'^www\.', '', url_without_protocol)
        text_without_protocol = re.sub(r'^https?://', '', text_cleaned)
        
        # Check if text is the same as URL (with or without protocol)
        if text_cleaned == url_cleaned:
            return True
        if text_without_protocol == url_without_protocol:
            return True
        if text_cleaned == url_without_www:
            return True
        
        return False

    def _is_related_heading(self, block) -> bool:
        text = self._norm(self._block_text(block))
        return any(rx.match(text) for rx in self.related_res)

    def _is_toc_line(self, text: str) -> bool:
        t = self._norm(text)
        if not t:
            return False
        return bool(self.dot_leader_re.search(t) or 
                   (self.page_number_trail_re.search(t) and t.count('.') >= 3))

    def _is_question_block(self, block) -> bool:
        """
        Identifies Q&A headings as either:
        1. Containing '?' (for questions)
        2. Starting with configurable problem/question patterns
        
        Uses configurable heading level range and skip_headings list.
        """
        if not isinstance(block, Paragraph):
            return False
        t = self._norm(block.text)
        if not t:
            return False
        
        # Skip headings that match skip_headings list (e.g., 'Summary')
        if t.lower() in self.skip_headings:
            return False
        
        lvl = self._heading_level(block)
        if lvl is None or lvl < self.min_heading_level or lvl > self.max_heading_level:
            return False
        
        # Accept if it has a question mark
        if '?' in t:
            return True
        
        # Accept configurable problem/question statement patterns
        t_lower = t.lower()
        return any(t_lower.startswith(pat) for pat in self.question_patterns)

    def _extract_hyperlinks_from_paragraph(self, paragraph) -> List[Dict[str, str]]:
        """Extract hyperlinks from a paragraph."""
        links = []
        if not isinstance(paragraph, Paragraph):
            return links
        
        try:
            part = paragraph.part
            rels = part.rels
        except Exception:
            return links
        
        for hyperlink in paragraph._element.xpath('.//w:hyperlink'):
            text_parts = []
            for run in hyperlink.xpath('.//w:t'):
                if run.text:
                    text_parts.append(run.text)
            text = ''.join(text_parts)
            
            if not text:
                continue
            
            # Try to get external URL via relationship ID
            r_id = hyperlink.get(qn('r:id'))
            if r_id and r_id in rels:
                url = rels[r_id].target_ref
                if url:
                    links.append({'text': text, 'url': url})
                    continue
            
            # Try to get internal anchor (bookmark link)
            # Skip scroll-bookmark internal links (e.g., #scroll-bookmark-17)
            anchor = hyperlink.get(qn('w:anchor'))
            if anchor and not anchor.startswith('scroll-bookmark'):
                links.append({'text': text, 'url': f'#{anchor}'})
        
        return links

    def _extract_hyperlinks_from_table(self, table) -> List[Dict[str, str]]:
        """Recursively extract hyperlinks from a table, including nested tables."""
        links = []
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    links.extend(self._extract_hyperlinks_from_paragraph(para))
                for nested_table in cell.tables:
                    links.extend(self._extract_hyperlinks_from_table(nested_table))
        return links

    def _extract_hyperlinks_from_block(self, block) -> List[Dict[str, str]]:
        """
        Extract hyperlinks from a block (paragraph or table).
        Returns list of dicts with 'text' and 'url' keys.
        """
        links = []
        
        if isinstance(block, Paragraph):
            links.extend(self._extract_hyperlinks_from_paragraph(block))
        elif isinstance(block, Table):
            links.extend(self._extract_hyperlinks_from_table(block))
        
        return links
